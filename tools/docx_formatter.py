"""ORCID-only DOCX formatter — generates a Word CV from ORCID JSON data.

Derived from tamu-coe-faculty-profiles/formatters/docx_formatter.py.
Contains only the ORCID-relevant renderers (orcid_data, publications)
and their helper functions.

See CLAUDE.md "Standalone CV Tool" section for upstream sync policy.

Requires: python-docx (pip install python-docx)
"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import logging

logger = logging.getLogger(__name__)

# Texas A&M maroon: RGB(80, 0, 0) — matches LaTeX \definecolor{tamumaroon}
TAMU_MAROON = RGBColor(80, 0, 0)


class OrcidDocxFormatter:
    """Generate a Word document from an ORCID-only faculty profile."""

    def format(self, profile: dict, output_dir: Path) -> list[Path]:
        """Generate a DOCX CV from the combined profile dict.

        Args:
            profile: Faculty profile dict with identity, orcid_data, orcid_publications
            output_dir: Directory for output files

        Returns:
            List containing the single DOCX file path
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        doc = Document()

        # Page setup: letter format (portrait), 1-inch margins
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # Style headings: TAMU maroon, left-aligned
        _setup_heading_styles(doc)

        _add_header(doc, profile["identity"], profile["_meta"])

        # Render ORCID sections in order
        render_order = [
            ("orcid_data", _render_orcid_data),
            ("orcid_publications", _render_publications),
        ]
        for key, renderer in render_order:
            data = profile.get(key)
            if not data:
                continue
            renderer(doc, data)

        report_id = profile["_meta"]["uin"]
        path = output_dir / f"{report_id}-cv.docx"
        doc.save(str(path))
        return [path]


# ---------------------------------------------------------------------------
# Header
# ---------------------------------------------------------------------------

def _add_header(doc, identity, meta):
    """Add title page matching the LaTeX report structure, with TOC."""
    # Institution name in maroon
    inst_para = doc.add_paragraph()
    inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_para.paragraph_format.space_after = Pt(4)
    run = inst_para.add_run(identity.get("institution", ""))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = TAMU_MAROON

    # College name
    college = identity.get("college", "")
    if college:
        college_para = doc.add_paragraph()
        college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        college_para.paragraph_format.space_after = Pt(12)
        run = college_para.add_run(college)
        run.font.size = Pt(14)

    # Horizontal rule
    _add_horizontal_rule(doc)

    # Faculty name (bold, large)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(14)
    name_para.paragraph_format.space_after = Pt(4)
    run = name_para.add_run(identity.get("name", ""))
    run.bold = True
    run.font.size = Pt(14)

    # Department name
    dept = identity.get("department_name", "")
    if dept:
        dept_para = doc.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_para.paragraph_format.space_after = Pt(4)
        run = dept_para.add_run(dept)
        run.font.size = Pt(12)

    # Identifier (ORCID)
    id_label = identity.get("identifier_label", "ORCID")
    id_value = identity.get("identifier_value", "")
    if id_value:
        id_para = doc.add_paragraph()
        id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        id_para.paragraph_format.space_after = Pt(10)
        run = id_para.add_run(f"{id_label}: {id_value}")
        run.font.size = Pt(9)

    # Horizontal rule
    _add_horizontal_rule(doc)

    # Year filter if present
    year_filter = meta.get("year_filter")
    if year_filter:
        year_para = doc.add_paragraph()
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_para.paragraph_format.space_before = Pt(10)
        run = year_para.add_run(f"Report Year: {year_filter}")
        run.font.size = Pt(11)

    # Spacing before TOC
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table of Contents
    _add_toc(doc)

    # Page break after title page
    doc.add_page_break()


# ---------------------------------------------------------------------------
# ORCID Data
# ---------------------------------------------------------------------------

def _render_orcid_data(doc, data):
    """Render ORCID profile data (biography, employment, education, etc.)."""
    doc.add_heading("ORCID Data", level=1)

    bio = data.get("biography")
    if bio:
        doc.add_paragraph(bio)

    employment = data.get("employment", [])
    if employment:
        doc.add_heading("Employment", level=2)
        headers = ["Position", "Organization", "Department", "Start", "End"]
        rows = [
            [
                e.get("role", ""),
                e.get("organization", ""),
                e.get("department", ""),
                e.get("start_year", ""),
                e.get("end_year", "") or "Present",
            ]
            for e in employment
        ]
        _add_table(doc, headers, rows,
                   col_widths=[1.5, 1.5, 1.5, 0.75, 0.75], font_size=9)

    education = data.get("education", [])
    if education:
        doc.add_heading("Education", level=2)
        headers = ["Degree", "Field of Study", "Organization", "Year"]
        rows = [
            [
                e.get("role", ""),
                e.get("department", ""),
                e.get("organization", ""),
                e.get("end_year", ""),
            ]
            for e in education
        ]
        _add_table(doc, headers, rows,
                   col_widths=[1, 2, 2, 1], font_size=9)

    # Remaining subsections as bulleted lists
    for key, label in [
        ("distinctions", "Distinctions"),
        ("memberships", "Memberships"),
        ("external_service", "External Service"),
        ("fundings", "Selected Projects"),
        ("external_identifiers", "External Identifiers"),
    ]:
        items = data.get(key, [])
        if not items:
            continue
        doc.add_heading(label, level=2)
        for item in items:
            parts = []
            role = item.get("role", "")
            org = item.get("organization", "")
            dept = item.get("department", "")
            title = item.get("title", "")
            # Primary label: role or title
            if role:
                parts.append(role)
            elif title:
                parts.append(title)
            # Organization (with department if present)
            if org and dept:
                parts.append(f"{org}, {dept}")
            elif org:
                parts.append(org)
            # Year range
            year_range = _fmt_year_range(
                item.get("start_year", ""), item.get("end_year", "")
            )
            if year_range:
                parts.append(year_range)
            # Fallback for items without standard fields (e.g., external_identifiers)
            if not parts:
                parts = [str(v) for v in item.values() if v]
            doc.add_paragraph(", ".join(parts), style="List Bullet")


# ---------------------------------------------------------------------------
# Publications
# ---------------------------------------------------------------------------

def _render_publications(doc, data):
    """Render publications (journal articles, conference papers, other)."""
    doc.add_heading("ORCID Publications", level=1)

    # Per-category counts
    categories = [
        ("journal_articles", "Journal Articles"),
        ("conference_papers", "Conference Papers"),
        ("other_publications", "Other Publications"),
    ]
    lines = []
    for key, label in categories:
        count = len(data.get(key, []))
        if count:
            lines.append(f"{count} {label} for the period considered")
    if lines:
        _add_summary_block(doc, lines)

    for key, label in categories:
        items = data.get(key, [])
        if not items:
            continue
        doc.add_heading(label, level=2)
        for i, pub in enumerate(items, 1):
            text = f"{i}. "
            if pub.get("authors"):
                text += pub["authors"] + ", "
            if pub.get("title"):
                text += f'"{pub["title"]}," '
            venue = (pub.get("venue") or pub.get("journal")
                     or pub.get("conference") or pub.get("publisher", ""))
            if venue:
                text += venue + ", "
            year = pub.get("year") or pub.get("publication_date", "")
            if year:
                text += f"{year}."
            para = doc.add_paragraph(text.rstrip(", "))
            doi = pub.get("doi")
            if doi:
                run = para.add_run()
                run._element.append(OxmlElement("w:br"))
                _add_hyperlink(para, f"DOI:{doi}", f"https://doi.org/{doi}")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_table(doc, headers, rows, *, col_widths=None, font_size=None):
    """Add a formatted table with a header row."""
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        table.autofit = False
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if font_size:
                    run.font.size = Pt(font_size)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            raw = str(value) if value is not None else ""
            cell.text = raw.replace("--", "-")
            if font_size:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
        if col_widths:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    # Spacing after table
    doc.add_paragraph()


def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run_el.append(rPr)
    run_el.text = text
    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


def _add_summary_block(doc, lines):
    """Add a compact summary block with line breaks (not paragraph spacing)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.first_line_indent = Pt(0)
    for i, line in enumerate(lines):
        if i > 0:
            run = para.add_run()
            run._element.append(OxmlElement("w:br"))
        para.add_run(line)


def _fmt_year_range(start_year: str, end_year: str) -> str:
    """Format a year range with en-dash: '2022\u20132025' or '2022\u2013present'."""
    if start_year and end_year:
        return f"{start_year}\u2013{end_year}"
    elif start_year:
        return f"{start_year}\u2013present"
    elif end_year:
        return end_year
    return ""


def _setup_heading_styles(doc):
    """Set all heading levels to TAMU maroon, left-aligned."""
    for level in range(1, 4):
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.color.rgb = TAMU_MAROON
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_horizontal_rule(doc):
    """Add a centered horizontal rule (60% page width) matching the LaTeX title page."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "500000")  # TAMU maroon hex
    pBdr.append(bottom)
    pPr.append(pBdr)
    # Indent to center the rule at ~60% of 6.5in page width
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(int(1.3 * 1440)))
    ind.set(qn("w:right"), str(int(1.3 * 1440)))
    pPr.append(ind)


def _add_toc(doc):
    """Insert a Table of Contents field (populated when opened in Word)."""
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = TAMU_MAROON

    para = doc.add_paragraph()
    # Begin field
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar)
    # Field instruction
    run = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._element.append(instrText)
    # Separate
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "separate")
    run._element.append(fldChar)
    # Placeholder text
    run = para.add_run("(Right-click and select 'Update Field' to populate)")
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(9)
    # End field
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar)


def _add_caption(doc, text):
    """Add an italic, centered caption paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
